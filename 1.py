# 导入处理Excel文件的库
import os
import sys
import json

# 清洗文件名中的非法字符
def clean_filename(filename):
    # Windows文件名非法字符
    illegal_chars = '\\/:*?"<>|'
    for char in illegal_chars:
        filename = filename.replace(char, '')
    # 移除不可见字符（换行符、制表符等）
    filename = filename.replace('\n', '').replace('\r', '').replace('\t', '')
    # 移除首尾空格
    filename = filename.strip()
    return filename
from math import ceil
from numpy.char import center
import pandas as pd
from openpyxl import Workbook, load_workbook

# 导入处理Word文件的库
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 导入tkinter模块创建弹窗
import tkinter as tk
from tkinter import messagebox, filedialog

# 导入上传文件的模块
import update_file

font_name = '仿宋_GB2312'
font_size = Pt(16)  # 三号字体约为16磅

# 配置文件路径
config_file = 'config.json'

# 默认保存路径
DEFAULT_SAVE_PATH = r'E:\共享文件\\4杜兴海\\6、综治信息平台纠纷附件\\2026矛盾纠纷附件'

# 加载配置
def load_config():
    if os.path.exists(config_file):
        with open(config_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

# 保存配置
def save_config(config):
    with open(config_file, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

# 创建文件选择和保存路径设置弹窗
def select_file_and_save_path():
    config = load_config()
    current_default = config.get('default_save_path', DEFAULT_SAVE_PATH)
    
    # 创建主窗口
    root = tk.Tk()
    root.title("选择文件和保存路径")
    root.geometry("600x300")
    root.resizable(False, False)
    
    # 存储选择的文件路径和保存路径
    selected_file = tk.StringVar()
    default_file = r'一级研判记录表.xlsx'
    selected_file.set(default_file)
    
    save_path_var = tk.StringVar(value=current_default)
    set_default_var = tk.BooleanVar(value=False)
    
    # 结果变量
    result = {'file_path': default_file, 'save_path': current_default, 'set_default': False, 'cancelled': True}
    
    def browse_file():
        file_path = filedialog.askopenfilename(
            title="选择Excel文件",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        if file_path:
            selected_file.set(file_path)
    
    def browse_save_path():
        selected = filedialog.askdirectory(
            title="选择保存路径",
            initialdir=current_default
        )
        if selected:
            save_path_var.set(selected)
    
    def confirm():
        result['file_path'] = selected_file.get()
        result['save_path'] = save_path_var.get()
        result['set_default'] = set_default_var.get()
        result['cancelled'] = False
        if result['set_default']:
            config['default_save_path'] = result['save_path']
            save_config(config)
        root.destroy()
    
    def cancel():
        result['cancelled'] = True
        root.destroy()
    
    # 处理窗口关闭事件
    def on_closing():
        result['cancelled'] = True
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # 创建界面元素
    tk.Label(root, text="请选择Excel文件：", font=("微软雅黑", 10)).pack(pady=10)
    
    file_frame = tk.Frame(root)
    file_frame.pack(pady=5)
    tk.Entry(file_frame, textvariable=selected_file, width=60).pack(side=tk.LEFT, padx=5)
    tk.Button(file_frame, text="浏览...", command=browse_file).pack(side=tk.LEFT)
    
    tk.Label(root, text="Word保存路径：", font=("微软雅黑", 10)).pack(pady=10)
    
    save_frame = tk.Frame(root)
    save_frame.pack(pady=5)
    tk.Entry(save_frame, textvariable=save_path_var, width=60).pack(side=tk.LEFT, padx=5)
    tk.Button(save_frame, text="浏览", command=browse_save_path).pack(side=tk.LEFT)
    
    # 设置默认复选框
    tk.Checkbutton(root, text="设为默认路径", variable=set_default_var).pack(pady=10)
    
    # 按钮区域
    button_frame = tk.Frame(root)
    button_frame.pack(pady=15)
    
    tk.Button(button_frame, text="确定", command=confirm, width=12, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=10)
    tk.Button(button_frame, text="取消", command=cancel, width=12).pack(side=tk.LEFT, padx=10)
    
    # 设置默认按钮为回车键
    root.bind('<Return>', lambda event: confirm())
    
    # 居中显示窗口
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    root.mainloop()
    return result

# 调用函数获取文件路径和保存路径
result = select_file_and_save_path()

# 检查是否取消
if result['cancelled']:
    sys.exit()

file_path = result['file_path']
save_base_path = result['save_path']

# 验证文件是否存在
if not os.path.exists(file_path):
    root = tk.Tk()
    root.withdraw()
    messagebox.showerror("错误", f"文件不存在：{file_path}")
    root.destroy()
    exit()

# 打开Excel文件地址
# file_path = r'一级研判记录表.xlsx'
# 定义设置中文字体的函数
def set_chinese_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 使用openpyxl读取Excel文件，保持原有格式
wb = load_workbook(file_path)
sheet = wb.active

# 获取表头
headers = [cell.value for cell in sheet[1]]

# 检查是否已存在链接列
if '链接' not in headers:
    # 添加链接列到最后
    sheet.cell(row=1, column=len(headers) + 1, value='链接')
    headers.append('链接')

# 遍历第二行开始的所有行
for row_idx in range(2, sheet.max_row + 1):
    # 获取当前行数据
    row_data = []
    for col_idx in range(1, len(headers) + 1):
        cell_value = sheet.cell(row=row_idx, column=col_idx).value
        row_data.append(cell_value)
    
    # 创建新的Word文档
    doc = Document()
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run('珉谷街道矛盾纠纷一级研判记录')
    title_run.font.name = '方正小标宋简体'
    # 设置中文字体
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    title_run.font.size = Pt(22)  # 二号字体约为22磅
    title.alignment = 1  # 居中对齐
 
    
    # 添加事件名称
    event_name_para = doc.add_paragraph()
    event_name_run = event_name_para.add_run(f'事件名称：{row_data[1] if len(row_data) > 1 else ""}')
    set_chinese_font(event_name_run, font_name)
    event_name_run.font.size = font_size  # 16号字体约为16磅
    event_name_para.alignment = 0  # 左对齐
    # 添加事件单号
    event_id_para = doc.add_paragraph()
    event_id_run = event_id_para.add_run(f'事件单号：{row_data[0] if len(row_data) > 0 else ""}')
    set_chinese_font(event_id_run, font_name)
    event_id_run.font.size = font_size  # 16号字体约为16磅
    event_id_para.alignment = 0  # 左对齐
# 添加事件描述标题
    event_textTitle_para = doc.add_paragraph()
    event_textTitle_run = event_textTitle_para.add_run('事件描述：')
    set_chinese_font(event_textTitle_run, font_name)
    event_textTitle_run.font.size = font_size  # 16号字体约为16磅
    event_textTitle_run.bold = True  # 加粗
    event_textTitle_para.alignment = 0  # 左对齐

# 添加事件描述
    event_text_para = doc.add_paragraph()
    event_text_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进两字符（约32磅）
    event_text_run = event_text_para.add_run(row_data[2] if len(row_data) > 2 else "")
    set_chinese_font(event_text_run, font_name)
    event_text_run.font.size = font_size  # 16号字体约为16磅
    event_text_para.alignment = 0  # 左对齐

# 添加风险研判标题
    event_riskAssessmentTitle_para = doc.add_paragraph()
    event_riskAssessmentTitle_run = event_riskAssessmentTitle_para.add_run('第一次研判及处置情况：')
    set_chinese_font(event_riskAssessmentTitle_run, font_name)
    event_riskAssessmentTitle_run.font.size = font_size  # 16号字体约为16磅
    event_riskAssessmentTitle_run.bold = True  # 加粗
    event_riskAssessmentTitle_para.alignment = 0  # 左对齐

# 添加风险研判
    event_riskAssessment_para = doc.add_paragraph()
    event_riskAssessment_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进两字符（约32磅）
    event_riskAssessment_run = event_riskAssessment_para.add_run(row_data[3] if len(row_data) > 3 else "")
    set_chinese_font(event_riskAssessment_run, font_name)
    event_riskAssessment_run.font.size = font_size  # 16号字体约为16磅
    event_riskAssessment_para.alignment = 0  # 左对齐

    
    # 设置文件保存路径和名称（使用行中的某个数据作为文件名，这里假设使用第一列数据）
    # 创建文件夹路径
    folder_name = clean_filename(str(row_data[0]) + '(' + str(row_data[1]) + ')' if len(row_data) > 1 else str(row_data[0]) if len(row_data) > 0 else "新建文件夹")
    folder_path = os.path.join(save_base_path, folder_name)
    # 如果文件夹不存在则创建
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    file_name = clean_filename("关于"+str(row_data[1]) + "的一级研判")
    save_path = os.path.join(folder_path, file_name + '.docx')
    # 保存文档
    try:
        doc.save(save_path)
        # 上传文件
        try:
            upload_result = update_file.upload_file(save_path)
            # 提取message
            message = upload_result.get('message', '')
            # 在Excel对应行的链接列添加内容
            sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=message)
        except Exception as upload_error:
            print(f"上传文件失败: {upload_error}")
            # 在Excel中记录上传失败信息
            sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=f"上传失败: {str(upload_error)}")
    except PermissionError:
        # 创建主窗口
        root = tk.Tk()
        # 隐藏主窗口
        root.withdraw()
        # 显示错误提示弹窗
        messagebox.showerror("权限错误", "请关闭《"+file_name+"》Word文档，然后重试。")
        # 关闭主窗口
        root.destroy()
    pass  # 在这里处理每一行的数据

# 保存Excel文件
try:
    wb.save(file_path)
except PermissionError:
    # 创建主窗口
    root = tk.Tk()
    # 隐藏主窗口
    root.withdraw()
    # 显示错误提示弹窗
    messagebox.showerror("权限错误", "请关闭Excel文件，然后重试。")
    # 关闭主窗口
    root.destroy()
    exit()

# 所有文件处理完成后显示成功提示
root = tk.Tk()
# 隐藏主窗口
root.withdraw()
# 显示成功提示弹窗
messagebox.showinfo("提示","运行成功")
# 关闭主窗口
root.destroy()

