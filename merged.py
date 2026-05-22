# 生成研判报告附件
# 导入处理Excel文件的库
import os
import sys
import json
import datetime
import threading
import queue

# 清洗文件名中的非法字符
def clean_filename(filename):
    # Windows文件名非法字符
    illegal_chars = '\\/:*?"<>|'
    for char in illegal_chars:
        filename = filename.replace(char, '')
    # 移除不可见字符（换行符、制表符等）
    filename = filename.replace('\n', '').replace('\r', '').replace('\t', '')
    # 移除首尾空格
    filename = filename.strip()
    return filename
from math import ceil
from numpy.char import center
import pandas as pd
from openpyxl import Workbook, load_workbook

# 导入处理Word文件的库
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 导入tkinter模块创建弹窗
import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk

# 配置文件路径
CONFIG_FILE = 'config.json'

# 导入上传文件的模块
import update_file

font_name = '仿宋_GB2312'
font_size = Pt(16)  # 三号字体约为16磅

# 配置文件路径
config_file = 'config.json'

# 默认保存路径
DEFAULT_SAVE_PATH = 'your_save_path_here'  # 请在config.json中配置default_save_path

# 加载配置
def load_config():
    if os.path.exists(config_file):
        with open(config_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

# 保存配置

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


def show_progress_bar(total, title="处理进度"):
    """
    显示进度条窗口
    :param total: 总任务数
    :param title: 窗口标题
    :return: 窗口对象、更新进度函数、关闭窗口函数
    """
    progress_window = tk.Toplevel()
    progress_window.title(title)
    progress_window.geometry("400x120")
    progress_window.resizable(False, False)
    progress_window.transient()
    
    progress_var = tk.DoubleVar()
    progress_bar = ttk.Progressbar(
        progress_window,
        variable=progress_var,
        maximum=100,
        length=350,
        mode='determinate'
    )
    progress_bar.pack(pady=20, padx=20)
    
    status_label = tk.Label(progress_window, text="准备开始...", font=("微软雅黑", 10))
    status_label.pack(pady=10)
    
    progress_window.update_idletasks()
    width = progress_window.winfo_width()
    height = progress_window.winfo_height()
    x = (progress_window.winfo_screenwidth() // 2) - (width // 2)
    y = (progress_window.winfo_screenheight() // 2) - (height // 2)
    progress_window.geometry(f'{width}x{height}+{x}+{y}')
    
    def update_progress(current, total, status):
        try:
            percent = (current / total) * 100
            progress_var.set(percent)
            status_label.config(text=status)
            progress_window.update()
        except tk.TclError:
            # 窗口已经被销毁，忽略错误
            pass
    
    def close():
        try:
            progress_window.destroy()
        except tk.TclError:
            # 窗口已经被销毁，忽略错误
            pass
    
    return progress_window, update_progress, close

# 创建文件选择和保存路径设置弹窗
def select_file_and_save_path():
    config = load_config()
    current_default = config.get('default_save_path', DEFAULT_SAVE_PATH)
    
    # 创建主窗口
    root = tk.Tk()
    root.title("选择文件和保存路径")
    root.geometry("700x300")
    root.resizable(False, False)
    
    # 存储选择的文件路径和保存路径
    selected_file = tk.StringVar()
    selected_file.set(r'一级研判记录表.xlsx')
    
    save_path_var = tk.StringVar(value=current_default)
    set_default_var = tk.BooleanVar(value=False)
    
    # 存储研判类型
    judgment_type = tk.StringVar(value="一级研判")
    
    # 存储token，自动加载已保存的token
    token_var = tk.StringVar(value=config.get('token', ''))
    
    # 结果变量
    result = {'file_path': selected_file.get(), 'save_path': current_default, 'set_default': False, 'cancelled': True, 'judgment_type': judgment_type.get(), 'token': ''}
    
    def browse_file():
        file_path = filedialog.askopenfilename(
            title="选择Excel文件",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        if file_path:
            selected_file.set(file_path)
    
    def browse_save_path():
        selected = filedialog.askdirectory(
            title="选择保存路径",
            initialdir=current_default
        )
        if selected:
            save_path_var.set(selected)
    
    def set_judgment_type(type_str):
        judgment_type.set(type_str)
        # 只改变研判类型，不改变Excel文件路径
    
    def confirm():
        result['file_path'] = selected_file.get()
        result['save_path'] = save_path_var.get()
        result['set_default'] = set_default_var.get()
        result['cancelled'] = False
        result['judgment_type'] = judgment_type.get()
        result['token'] = token_var.get()
        if result['set_default']:
            config['default_save_path'] = result['save_path']
        # 保存token
        if result['token']:
            config['token'] = result['token']
        save_config(config)
        root.destroy()
    
    def cancel():
        result['cancelled'] = True
        root.destroy()
    
    # 处理窗口关闭事件
    def on_closing():
        result['cancelled'] = True
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # 创建界面元素
    # Excel文件选择
    file_frame = tk.Frame(root)
    file_frame.pack(pady=10, fill=tk.X, padx=20)
    tk.Label(file_frame, text="请选择Excel文件：", font=("微软雅黑", 10), width=15, anchor=tk.W).pack(side=tk.LEFT, padx=5)
    entry_frame = tk.Frame(file_frame)
    entry_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
    tk.Entry(entry_frame, textvariable=selected_file, width=50).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    tk.Button(entry_frame, text="浏览...", command=browse_file, relief=tk.FLAT, bg="#f0f0f0", padx=10, pady=2).pack(side=tk.LEFT)
    
    # Word保存路径
    save_frame = tk.Frame(root)
    save_frame.pack(pady=10, fill=tk.X, padx=20)
    tk.Label(save_frame, text="Word保存路径：", font=("微软雅黑", 10), width=15, anchor=tk.W).pack(side=tk.LEFT, padx=5)
    save_entry_frame = tk.Frame(save_frame)
    save_entry_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
    tk.Entry(save_entry_frame, textvariable=save_path_var, width=50).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    tk.Button(save_entry_frame, text="浏览", command=browse_save_path, relief=tk.FLAT, bg="#f0f0f0", padx=10, pady=2).pack(side=tk.LEFT, padx=5)
    tk.Checkbutton(save_entry_frame, text="默认", variable=set_default_var).pack(side=tk.LEFT, padx=5)
    
    # Token输入框
    token_frame = tk.Frame(root)
    token_frame.pack(pady=10, fill=tk.X, padx=20)
    tk.Label(token_frame, text="用户令牌：", font=("微软雅黑", 10), width=15, anchor=tk.W).pack(side=tk.LEFT, padx=5)
    tk.Entry(token_frame, textvariable=token_var, width=50).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    
    # 研判类型选择
    type_frame = tk.Frame(root)
    type_frame.pack(pady=10, fill=tk.X, padx=20)
    tk.Label(type_frame, text="研判类型：", font=("微软雅黑", 10), width=15, anchor=tk.W).pack(side=tk.LEFT, padx=5)
    radio_frame = tk.Frame(type_frame)
    radio_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
    tk.Radiobutton(radio_frame, text="一级研判", variable=judgment_type, value="一级研判", command=lambda: set_judgment_type("一级研判")).pack(side=tk.LEFT, padx=10)
    tk.Radiobutton(radio_frame, text="二级研判", variable=judgment_type, value="二级研判", command=lambda: set_judgment_type("二级研判")).pack(side=tk.LEFT, padx=10)
    tk.Radiobutton(radio_frame, text="周研判", variable=judgment_type, value="周研判", command=lambda: set_judgment_type("周研判")).pack(side=tk.LEFT, padx=10)
    
    # 按钮区域
    button_frame = tk.Frame(root)
    button_frame.pack(pady=15)
    
    tk.Button(button_frame, text="确定", command=confirm, width=12, bg="#4CAF50", fg="white", relief=tk.FLAT, padx=15, pady=5).pack(side=tk.LEFT, padx=10)
    tk.Button(button_frame, text="取消", command=cancel, width=12, relief=tk.FLAT, bg="#f0f0f0", padx=15, pady=5).pack(side=tk.LEFT, padx=10)
    
    # 设置默认按钮为回车键
    root.bind('<Return>', lambda event: confirm())
    
    # 居中显示窗口
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    root.mainloop()
    return result

# 调用函数获取文件路径和保存路径
result = select_file_and_save_path()

# 检查是否取消
if result['cancelled']:
    sys.exit()

file_path = result['file_path']
save_base_path = result['save_path']
judgment_type = result['judgment_type']
token = result['token']

# 验证文件是否存在
if not os.path.exists(file_path):
    root = tk.Tk()
    root.withdraw()
    messagebox.showerror("错误", f"文件不存在：{file_path}")
    root.destroy()
    exit()

# 定义设置中文字体的函数
def set_chinese_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 使用openpyxl读取Excel文件，保持原有格式
wb = load_workbook(file_path)
sheet = wb.active

# 获取表头
headers = [cell.value for cell in sheet[1]]

# 检查是否已存在链接列
if '链接' not in headers:
    # 添加链接列到最后
    sheet.cell(row=1, column=len(headers) + 1, value='链接')
    headers.append('链接')

# 获取当前日期，格式：yyyy年mm月dd日
today = datetime.datetime.now()
date_str = today.strftime('%Y年%m月%d日')

# 获取总行数
total_rows = sheet.max_row - 1  # 减去表头

# 创建主窗口用于进度条
root = tk.Tk()
root.withdraw()

# 全局变量
progress_window = None
update_progress = None
close_progress = None

# 创建消息队列
msg_queue = queue.Queue()

# 显示进度条
progress_window, update_progress, close_progress = show_progress_bar(total_rows)

# 处理队列中的消息
def process_queue():
    global progress_window, update_progress, close_progress
    while not msg_queue.empty():
        try:
            msg = msg_queue.get_nowait()
            if msg['type'] == 'close_progress':
                # 关闭进度条
                if progress_window:
                    progress_window.destroy()
            elif msg['type'] == 'show_token_dialog':
                # 弹出令牌输入窗口
                new_token = tk.StringVar(value=token)
                token_window = tk.Toplevel()
                token_window.title("权限不足")
                token_window.geometry("500x200")
                token_window.resizable(False, False)
                token_window.grab_set()
                
                def confirm_token():
                    msg['callback'](new_token.get(), False)
                    token_window.destroy()
                
                def cancel_token():
                    msg['callback']("", True)
                    token_window.destroy()
                
                tk.Label(token_window, text="权限不足，请重新输入用户令牌：", font=("微软雅黑", 10)).pack(pady=20)
                
                token_frame = tk.Frame(token_window)
                token_frame.pack(pady=10)
                tk.Entry(token_frame, textvariable=new_token, width=50).pack(side=tk.LEFT, padx=5)
                
                button_frame = tk.Frame(token_window)
                button_frame.pack(pady=20)
                
                tk.Button(button_frame, text="确定", command=confirm_token, width=12, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=10)
                tk.Button(button_frame, text="取消", command=cancel_token, width=12).pack(side=tk.LEFT, padx=10)
                
                # 居中显示窗口
                token_window.update_idletasks()
                width = token_window.winfo_width()
                height = token_window.winfo_height()
                x = (token_window.winfo_screenwidth() // 2) - (width // 2)
                y = (token_window.winfo_screenheight() // 2) - (height // 2)
                token_window.geometry(f'{width}x{height}+{x}+{y}')
            elif msg['type'] == 'show_progress':
                # 重新显示进度条
                progress_window, update_progress, close_progress = show_progress_bar(total_rows)
                update_progress(msg['current_row'] - 1, total_rows, f"处理第 {msg['current_row']} 行: {msg['row_data']}")
            elif msg['type'] == 'show_error_and_exit':
                # 关闭进度条
                if progress_window:
                    try:
                        progress_window.destroy()
                    except tk.TclError:
                        pass
                
                # 显示错误提示并退出程序
                error_window = tk.Toplevel()
                error_window.title("权限不足")
                error_window.geometry("500x200")
                error_window.resizable(False, False)
                error_window.grab_set()
                error_window.attributes('-topmost', True)
                
                tk.Label(error_window, text="权限不足，请重新输入令牌并重启程序。", font=("微软雅黑", 10)).pack(pady=30)
                
                button_frame = tk.Frame(error_window)
                button_frame.pack(pady=20)
                
                def exit_program():
                    error_window.destroy()
                    # 保存Excel文件
                    try:
                        wb.save(file_path)
                    except Exception:
                        pass
                    # 强制退出程序
                    os._exit(1)
                
                tk.Button(button_frame, text="确定", command=exit_program, width=12, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=10)
                
                # 居中显示窗口
                error_window.update_idletasks()
                width = error_window.winfo_width()
                height = error_window.winfo_height()
                x = (error_window.winfo_screenwidth() // 2) - (width // 2)
                y = (error_window.winfo_screenheight() // 2) - (height // 2)
                error_window.geometry(f'{width}x{height}+{x}+{y}')
                
                # 立即处理消息，确保错误窗口显示
                error_window.update()
        except queue.Empty:
            break
    # 继续处理队列
    root.after(100, process_queue)

# 启动队列处理
root.after(100, process_queue)

try:
    # 遍历第二行开始的所有行
    for row_idx in range(2, sheet.max_row + 1):
        current_row = row_idx - 1  # 当前处理的行（从1开始）
        # 获取当前行数据
        row_data = []
        for col_idx in range(1, len(headers) + 1):
            cell_value = sheet.cell(row=row_idx, column=col_idx).value
            row_data.append(cell_value)
        # 更新进度条
        update_progress(current_row - 1, total_rows, f"处理第 {current_row} 行: {row_data[1] if len(row_data) > 1 else '未知'}")
        
        # 创建新的Word文档
        doc = Document()
        # 添加标题
        title = doc.add_paragraph()
        
        # 根据研判类型设置标题
        if judgment_type == "一级研判":
            title_text = '珉谷街道矛盾纠纷一级研判记录'
        elif judgment_type == "二级研判":
            title_text = '珉谷街道矛盾纠纷二级研判记录'
        elif judgment_type == "周研判":
            title_text = '珉谷街道矛盾纠纷周研判记录'
        
        title_run = title.add_run(title_text)
        title_run.font.name = '方正小标宋简体'
        # 设置中文字体
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        # 设置字体大小为22号
        title_run.font.size = Pt(22)
        # 居中对齐
        title.alignment = 1  # 居中对齐

        # 添加事件名称
        event_name_para = doc.add_paragraph()
        event_name_run = event_name_para.add_run(f'事件名称：{row_data[1] if len(row_data) > 1 else ""}')
        set_chinese_font(event_name_run, font_name)
        event_name_run.font.size = font_size  # 16号字体约为16磅
        event_name_para.alignment = 0  # 左对齐
        # 添加事件单号
        event_id_para = doc.add_paragraph()
        event_id_run = event_id_para.add_run(f'事件单号：{row_data[0] if len(row_data) > 0 else ""}')
        set_chinese_font(event_id_run, font_name)
        event_id_run.font.size = font_size  # 16号字体约为16磅
        event_id_para.alignment = 0  # 左对齐
        # 添加事件描述标题
        event_textTitle_para = doc.add_paragraph()
        event_textTitle_run = event_textTitle_para.add_run('事件描述：')
        set_chinese_font(event_textTitle_run, font_name)
        event_textTitle_run.font.size = font_size  # 16号字体约为16磅
        event_textTitle_run.bold = True  # 加粗
        event_textTitle_para.alignment = 0  # 左对齐

        # 添加事件描述
        event_text_para = doc.add_paragraph()
        event_text_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进两字符（约32磅）
        event_text_run = event_text_para.add_run(row_data[2] if len(row_data) > 2 else "")
        set_chinese_font(event_text_run, font_name)
        event_text_run.font.size = font_size  # 16号字体约为16磅
        event_text_para.alignment = 0  # 左对齐

        # 添加风险研判标题
        event_riskAssessmentTitle_para = doc.add_paragraph()
        event_riskAssessmentTitle_run = event_riskAssessmentTitle_para.add_run('第一次研判及处置情况：')
        set_chinese_font(event_riskAssessmentTitle_run, font_name)
        event_riskAssessmentTitle_run.font.size = font_size  # 16号字体约为16磅
        event_riskAssessmentTitle_run.bold = True  # 加粗
        event_riskAssessmentTitle_para.alignment = 0  # 左对齐

        # 添加风险研判
        event_riskAssessment_para = doc.add_paragraph()
        event_riskAssessment_para.paragraph_format.first_line_indent = Pt(32)  # 首行缩进两字符（约32磅）
        event_riskAssessment_run = event_riskAssessment_para.add_run(row_data[3] if len(row_data) > 3 else "")
        set_chinese_font(event_riskAssessment_run, font_name)
        event_riskAssessment_run.font.size = font_size  # 16号字体约为16磅
        event_riskAssessment_para.alignment = 0  # 左对齐

        # 设置文件保存路径和名称（使用行中的某个数据作为文件名，这里假设使用第一列数据）
        # 创建文件夹路径
        folder_name = clean_filename(str(row_data[0]) + '(' + str(row_data[1]) + ')' if len(row_data) > 1 else str(row_data[0]) if len(row_data) > 0 else "新建文件夹")
        folder_path = os.path.join(save_base_path, folder_name)
        # 如果文件夹不存在则创建
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        # 根据研判类型设置文件名
        if judgment_type == "一级研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的一级研判")
        elif judgment_type == "二级研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的二级研判")
        elif judgment_type == "周研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的周研判")
        
        # 添加日期到文件名
        file_name_with_date = f"{file_name}_{date_str}"
        save_path = os.path.join(folder_path, file_name_with_date + '.docx')
        
        # 保存文档
        try:
            doc.save(save_path)
            # 上传文件
            def upload_task():
                global token
                try:
                    upload_result = update_file.upload_file(save_path, token)
                    # 提取message
                    message = upload_result.get('message', '')
                    # 在Excel对应行的链接列添加内容
                    sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=message)
                except Exception as upload_error:
                    error_msg = str(upload_error)
                    if "401" in error_msg or "权限不足" in error_msg or "unauthorized" in error_msg.lower():
                        # 关闭进度条
                        msg_queue.put({'type': 'close_progress'})
                        # 显示权限不足提示并停止程序
                        def show_error_and_exit():
                            # 弹出错误提示
                            error_window = tk.Toplevel()
                            error_window.title("权限不足")
                            error_window.geometry("500x200")
                            error_window.resizable(False, False)
                            error_window.grab_set()
                            
                            tk.Label(error_window, text="权限不足，请重新输入令牌并重启程序。", font=("微软雅黑", 10)).pack(pady=30)
                            
                            button_frame = tk.Frame(error_window)
                            button_frame.pack(pady=20)
                            
                            def exit_program():
                                error_window.destroy()
                                # 保存Excel文件
                                try:
                                    wb.save(file_path)
                                except Exception:
                                    pass
                                # 退出程序
                                sys.exit()
                            
                            tk.Button(button_frame, text="确定", command=exit_program, width=12, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=10)
                            
                            # 居中显示窗口
                            error_window.update_idletasks()
                            width = error_window.winfo_width()
                            height = error_window.winfo_height()
                            x = (error_window.winfo_screenwidth() // 2) - (width // 2)
                            y = (error_window.winfo_screenheight() // 2) - (height // 2)
                            error_window.geometry(f'{width}x{height}+{x}+{y}')
                        
                        # 显示错误提示并退出
                        msg_queue.put({
                            'type': 'show_error_and_exit'
                        })
                    else:
                        print(f"上传文件失败: {upload_error}")
                        # 在Excel中记录上传失败信息
                        sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=f"上传失败: {str(upload_error)}")
            
            # 创建并启动上传线程
            upload_thread = threading.Thread(target=upload_task)
            upload_thread.daemon = True
            upload_thread.start()
            
            # 等待上传线程完成
            upload_thread.join()
        except PermissionError:
            # 创建主窗口
            root = tk.Tk()
            # 隐藏主窗口
            root.withdraw()
            # 显示错误提示弹窗
            messagebox.showerror("权限错误", "请关闭《"+file_name_with_date+"》Word文档，然后重试。")
            # 关闭主窗口
            root.destroy()
        pass  # 在这里处理每一行的数据

finally:
    # 关闭进度条
    update_progress(total_rows, total_rows, "任务完成！")
    import time
    time.sleep(0.5)
    close_progress()

# 保存Excel文件
try:
    wb.save(file_path)
except PermissionError:
    # 创建主窗口
    root = tk.Tk()
    # 隐藏主窗口
    root.withdraw()
    # 显示错误提示弹窗
    messagebox.showerror("权限错误", "请关闭Excel文件，然后重试。")
    # 关闭主窗口
    root.destroy()
    exit()

# 所有文件处理完成后显示成功提示
root = tk.Tk()
# 隐藏主窗口
root.withdraw()
# 显示成功提示弹窗
messagebox.showinfo("提示","运行成功")
# 关闭主窗口
root.destroy()