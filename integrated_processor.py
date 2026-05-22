import os
import sys
import json
import datetime
import requests
import urllib3
import pandas as pd
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import ttk

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

CONFIG_FILE = 'config.json'
DATA_FILE = 'dispute_data.json'
DEFAULT_SAVE_PATH = r'E:\共享文件\\4杜兴海\\6、综治信息平台纠纷附件\\2026矛盾纠纷附件'
font_name = '仿宋_GB2312'
font_size = Pt(16)

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def get_base_url():
    """从配置文件获取服务器基础URL"""
    config = load_config()
    host = config.get('server_host', 'your_server_ip_here')
    port = config.get('server_port', 'your_server_port_here')
    return f'https://{host}:{port}'

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

def clean_filename(filename):
    illegal_chars = '\\/:*?"<>|'
    for char in illegal_chars:
        filename = filename.replace(char, '')
    filename = filename.replace('\n', '').replace('\r', '').replace('\t', '')
    return filename.strip()

def get_new_token_dialog(current_token=""):
    root = tk._default_root
    if root is None:
        root = tk.Tk()
        root.withdraw()
    
    new_token = tk.StringVar(value=current_token)
    token_window = tk.Toplevel(root)
    token_window.title("权限不足")
    token_window.geometry("500x200")
    token_window.resizable(False, False)
    token_window.grab_set()
    
    result = {"token": "", "cancelled": True}
    
    def confirm_token():
        result["token"] = new_token.get()
        result["cancelled"] = False
        token_window.destroy()
    
    def cancel_token():
        result["cancelled"] = True
        token_window.destroy()
    
    tk.Label(token_window, text="权限不足，请重新输入用户令牌：", font=("微软雅黑", 10)).pack(pady=20)
    
    token_frame = tk.Frame(token_window)
    token_frame.pack(pady=10)
    tk.Entry(token_frame, textvariable=new_token, width=50).pack(side=tk.LEFT, padx=5)
    
    button_frame = tk.Frame(token_window)
    button_frame.pack(pady=20)
    
    tk.Button(button_frame, text="确定", command=confirm_token, width=12, bg="#4CAF50", fg="white").pack(side=tk.LEFT, padx=10)
    tk.Button(button_frame, text="取消", command=cancel_token, width=12).pack(side=tk.LEFT, padx=10)
    
    token_window.update_idletasks()
    width = token_window.winfo_width()
    height = token_window.winfo_height()
    x = (token_window.winfo_screenwidth() // 2) - (width // 2)
    y = (token_window.winfo_screenheight() // 2) - (height // 2)
    token_window.geometry(f'{width}x{height}+{x}+{y}')
    
    token_window.wait_window()
    return result

def upload_file(file_path, token):
    url = f'{get_base_url()}/sys/common/upload'
    headers = {'x-access-token': token}
    
    try:
        with open(file_path, 'rb') as f:
            files = {'file': f}
            response = requests.post(url, headers=headers, files=files, verify=False, timeout=30)
        
        response.raise_for_status()
        return response.json()
    except Exception as error:
        print(f'上传错误: {error}')
        raise

def fetch_dispute_data(token, update_progress_callback=None):
    if update_progress_callback:
        update_progress_callback(0, 100, "步骤1/4: 从API获取矛盾纠纷数据...")
    
    base_url = get_base_url()
    urls = [
        ("矛盾纠纷", f"{base_url}/event-center/homePage/getMyTodoList?_t={datetime.datetime.now().timestamp()}&pageNo=1&pageSize=1000&sort=createTime&order=descend&serialNumber=&subject="),
        ("贵和码", f"{base_url}/event-center/homePage/getMyTodoList?_t={datetime.datetime.now().timestamp()}&pageNo=1&pageSize=1000&pageType=2&sort=createTime&order=descend&serialNumber=&subject="),
        ("矛盾纠纷登记", f"{base_url}/event-center/homePage/getMyTodoList?_t={datetime.datetime.now().timestamp()}&pageNo=1&pageSize=1000&pageType=1&sort=createTime&order=descend&serialNumber=&subject=")
    ]
    
    all_data = []
    
    for i, (data_type, url) in enumerate(urls):
        try:
            headers = {'x-access-token': token}
            response = requests.get(url, headers=headers, verify=False)
            response.raise_for_status()
            result = response.json()
            
            if 'result' in result and 'records' in result['result']:
                records = result['result']['records']
                for record in records:
                    all_data.append({
                        'id': record.get('id'),
                        'circulationId': record.get('circulationId'),
                        'serialNumber': record.get('serialNumber'),
                        'name': record.get('name'),
                        'createTime': record.get('createTime'),
                        'remainTime': record.get('remainTime'),
                        'riskLevel': record.get('riskLevel'),
                        'dataType': data_type
                    })
            
            if update_progress_callback:
                progress = 25 * (i + 1) / len(urls)
                update_progress_callback(progress, 100, f"步骤1/4: 获取{data_type}数据成功，共{len(records)}条")
            
            print(f"获取{data_type}数据成功，共{len(records)}条")
        except Exception as e:
            print(f"获取{data_type}数据失败: {e}")
    
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)
    
    if update_progress_callback:
        update_progress_callback(25, 100, f"步骤1/4: 数据已保存到 {DATA_FILE}，共 {len(all_data)} 条记录")
    
    print(f"数据已保存到 {DATA_FILE}，共 {len(all_data)} 条记录")
    return all_data

def find_record_by_serial_number(serial_number):
    if not os.path.exists(DATA_FILE):
        return None
    
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    for record in data:
        if str(record.get('serialNumber', '')).strip() == str(serial_number).strip():
            return record
    return None

def update_progress_api(issue_id, circulation_id, opinion, token, attachment_list=None):
    url = f'{get_base_url()}/event-center/api/circulation/updateProgress'
    headers = {
        'x-access-token': token,
        'Content-Type': 'application/json'
    }
    
    data = {
        "issueId": issue_id,
        "circulationId": circulation_id,
        "operationType": 100009,
        "mediationType": 2,
        "opinion": opinion
    }
    
    if attachment_list:
        data["attachmentList"] = attachment_list
    
    response = requests.post(url, headers=headers, json=data, verify=False)
    response.raise_for_status()
    return response.json()

def set_chinese_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def generate_judgment_reports(excel_path, save_path, judgment_type, token, update_progress_callback=None):
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"文件不存在：{excel_path}")
    
    wb = load_workbook(excel_path)
    sheet = wb.active
    headers = [cell.value for cell in sheet[1]]
    
    if '链接' not in headers:
        sheet.cell(row=1, column=len(headers) + 1, value='链接')
        headers.append('链接')
    
    today = datetime.datetime.now()
    date_str = today.strftime('%Y年%m月%d日')
    total_rows = sheet.max_row - 1
    
    if update_progress_callback:
        update_progress_callback(25, 100, "步骤2/4: 开始生成研判报告...")
    
    for row_idx in range(2, sheet.max_row + 1):
        current_row = row_idx - 1
        
        if update_progress_callback:
            progress = 25 + (current_row / total_rows) * 25
            update_progress_callback(progress, 100, f"步骤2/4: 生成报告 {current_row}/{total_rows}")
        
        row_data = []
        for col_idx in range(1, len(headers) + 1):
            row_data.append(sheet.cell(row=row_idx, column=col_idx).value)
        
        doc = Document()
        title = doc.add_paragraph()
        
        if judgment_type == "一级研判":
            title_text = '珉谷街道矛盾纠纷一级研判记录'
        elif judgment_type == "二级研判":
            title_text = '珉谷街道矛盾纠纷二级研判记录'
        elif judgment_type == "周研判":
            title_text = '珉谷街道矛盾纠纷周研判记录'
        
        title_run = title.add_run(title_text)
        title_run.font.name = '方正小标宋简体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
        title_run.font.size = Pt(22)
        title.alignment = 1
        
        event_name_para = doc.add_paragraph()
        event_name_run = event_name_para.add_run(f'事件名称：{row_data[1] if len(row_data) > 1 else ""}')
        set_chinese_font(event_name_run, font_name)
        event_name_run.font.size = font_size
        event_name_para.alignment = 0
        
        event_id_para = doc.add_paragraph()
        event_id_run = event_id_para.add_run(f'事件单号：{row_data[0] if len(row_data) > 0 else ""}')
        set_chinese_font(event_id_run, font_name)
        event_id_run.font.size = font_size
        event_id_para.alignment = 0
        
        event_textTitle_para = doc.add_paragraph()
        event_textTitle_run = event_textTitle_para.add_run('事件描述：')
        set_chinese_font(event_textTitle_run, font_name)
        event_textTitle_run.font.size = font_size
        event_textTitle_run.bold = True
        event_textTitle_para.alignment = 0
        
        event_text_para = doc.add_paragraph()
        event_text_para.paragraph_format.first_line_indent = Pt(32)
        event_text_run = event_text_para.add_run(row_data[2] if len(row_data) > 2 else "")
        set_chinese_font(event_text_run, font_name)
        event_text_run.font.size = font_size
        event_text_para.alignment = 0
        
        event_riskAssessmentTitle_para = doc.add_paragraph()
        event_riskAssessmentTitle_run = event_riskAssessmentTitle_para.add_run('第一次研判及处置情况：')
        set_chinese_font(event_riskAssessmentTitle_run, font_name)
        event_riskAssessmentTitle_run.font.size = font_size
        event_riskAssessmentTitle_run.bold = True
        event_riskAssessmentTitle_para.alignment = 0
        
        event_riskAssessment_para = doc.add_paragraph()
        event_riskAssessment_para.paragraph_format.first_line_indent = Pt(32)
        event_riskAssessment_run = event_riskAssessment_para.add_run(row_data[3] if len(row_data) > 3 else "")
        set_chinese_font(event_riskAssessment_run, font_name)
        event_riskAssessment_run.font.size = font_size
        event_riskAssessment_para.alignment = 0
        
        folder_name = clean_filename(str(row_data[0]) + '(' + str(row_data[1]) + ')' if len(row_data) > 1 else str(row_data[0]) if len(row_data) > 0 else "新建文件夹")
        folder_path = os.path.join(save_path, folder_name)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        if judgment_type == "一级研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的一级研判")
        elif judgment_type == "二级研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的二级研判")
        elif judgment_type == "周研判":
            file_name = clean_filename("关于"+str(row_data[1]) + "的周研判")
        
        file_name_with_date = f"{file_name}_{date_str}"
        doc_save_path = os.path.join(folder_path, file_name_with_date + '.docx')
        
        doc.save(doc_save_path)
        
        try:
            upload_result = upload_file(doc_save_path, token)
            message = upload_result.get('message', '')
            sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=message)
            print(f"上传成功: {file_name_with_date}")
        except Exception as upload_error:
            error_msg = str(upload_error)
            if "401" in error_msg or "权限不足" in error_msg or "unauthorized" in error_msg.lower():
                token_result = get_new_token_dialog(token)
                if token_result["cancelled"]:
                    raise Exception("用户取消输入令牌")
                token = token_result["token"]
                config = load_config()
                config['token'] = token
                save_config(config)
                upload_result = upload_file(doc_save_path, token)
                message = upload_result.get('message', '')
                sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=message)
            else:
                sheet.cell(row=row_idx, column=headers.index('链接') + 1, value=f"上传失败: {str(upload_error)}")
    
    try:
        wb.save(excel_path)
    except PermissionError:
        if tk._default_root:
            messagebox.showerror("权限错误", "请关闭Excel文件，然后重试。")
        else:
            temp_root = tk.Tk()
            temp_root.withdraw()
            messagebox.showerror("权限错误", "请关闭Excel文件，然后重试。")
            temp_root.destroy()
        exit()
    
    if update_progress_callback:
        update_progress_callback(50, 100, "步骤2/4: 研判报告生成并上传完成")
    
    return token

def process_update_progress(judgment_excel_path, token, update_progress_callback=None):
    if update_progress_callback:
        update_progress_callback(50, 100, "步骤3/4: 查找id和circulationId...")
    
    wb = load_workbook(judgment_excel_path)
    sheet = wb.active
    headers = [cell.value for cell in sheet[1]]
    
    id_col = None
    circulation_id_col = None
    remark_col = None
    
    for col_idx, header in enumerate(headers):
        if header and str(header).strip() == 'id':
            id_col = col_idx
        elif header and str(header).strip() == 'circulationId':
            circulation_id_col = col_idx
        elif header and str(header).strip() == '上传备注':
            remark_col = col_idx
    
    if id_col is None:
        id_col = len(headers)
        sheet.cell(row=1, column=id_col + 1, value='id')
        headers.append('id')
    
    if circulation_id_col is None:
        circulation_id_col = len(headers)
        sheet.cell(row=1, column=circulation_id_col + 1, value='circulationId')
        headers.append('circulationId')
    
    if remark_col is None:
        remark_col = len(headers)
        sheet.cell(row=1, column=remark_col + 1, value='上传备注')
        headers.append('上传备注')
    
    order_col = 0
    opinion_col = None
    link_col = None
    
    for col_idx, col_name in enumerate(headers):
        if col_name:
            col_name_str = str(col_name).lower()
            if '研判' in col_name_str or '处置' in col_name_str or '意见' in col_name_str:
                opinion_col = col_idx
            elif '链接' in col_name_str or 'url' in col_name_str or 'message' in col_name_str:
                link_col = col_idx
    
    if opinion_col is None:
        opinion_col = 3
    if link_col is None:
        link_col = 4
    
    records = []
    total_rows = sheet.max_row - 1
    
    skipped_count = 0
    
    for row_idx in range(2, sheet.max_row + 1):
        order_no_cell = sheet.cell(row=row_idx, column=order_col + 1)
        if order_no_cell.value:
            order_no = str(order_no_cell.value).strip()
            if order_no:
                opinion_cell = sheet.cell(row=row_idx, column=opinion_col + 1) if opinion_col + 1 <= len(headers) else None
                link_cell = sheet.cell(row=row_idx, column=link_col + 1) if link_col + 1 <= len(headers) else None
                
                opinion = str(opinion_cell.value) if opinion_cell and opinion_cell.value else ''
                url = str(link_cell.value) if link_cell and link_cell.value else ''
                
                record_from_data = find_record_by_serial_number(order_no)
                if not record_from_data:
                    print(f"单号 {order_no} 未找到对应数据，跳过")
                    skipped_count += 1
                    continue
                
                issue_id = record_from_data['id']
                circulation_id = record_from_data['circulationId']
                
                sheet.cell(row=row_idx, column=id_col + 1, value=issue_id)
                sheet.cell(row=row_idx, column=circulation_id_col + 1, value=circulation_id)
                
                records.append({
                    'row_idx': row_idx,
                    'order_no': order_no,
                    'opinion': opinion,
                    'link': url,
                    'issue_id': issue_id,
                    'circulation_id': circulation_id,
                    'remark_col': remark_col
                })
    
    if skipped_count > 0:
        print(f"共跳过 {skipped_count} 条未找到对应单号的记录")
    
    try:
        wb.save(judgment_excel_path)
        print(f"已将id和circulationId写入研判记录表")
    except PermissionError:
        if tk._default_root:
            messagebox.showerror("权限错误", "请关闭研判记录表Excel文件，然后重试。")
        else:
            temp_root = tk.Tk()
            temp_root.withdraw()
            messagebox.showerror("权限错误", "请关闭研判记录表Excel文件，然后重试。")
            temp_root.destroy()
        return token
    
    if update_progress_callback:
        update_progress_callback(60, 100, f"步骤3/4: id和circulationId已写入，共 {len(records)} 条记录")
    
    print(f"读取研判记录表成功，共 {len(records)} 条记录")
    
    success_count = 0
    failure_count = 0
    failed_order_nos = []
    
    if update_progress_callback:
        update_progress_callback(60, 100, "步骤4/4: 开始上传研判记录...")
    
    for idx, record in enumerate(records):
        current_row = idx + 1
        order_no = record['order_no']
        opinion = record['opinion']
        url = record['link']
        issue_id = record['issue_id']
        circulation_id = record['circulation_id']
        row_idx = record['row_idx']
        remark_col = record['remark_col']
        
        if update_progress_callback:
            progress = 60 + (current_row / len(records)) * 40
            update_progress_callback(progress, 100, f"步骤4/4: 上传记录 {current_row}/{len(records)}")
        
        print(f"处理第 {current_row} 行: issueId={issue_id}, circulationId={circulation_id}, orderNo={order_no}")
        
        attachment_list = None
        if url:
            ext = os.path.splitext(url)[1].lower()
            mime_types = {
                '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
                '.gif': 'image/gif', '.bmp': 'image/bmp', '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.pdf': 'application/pdf', '.txt': 'text/plain',
                '.xls': 'application/vnd.ms-excel',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
            file_type = mime_types.get(ext, 'application/octet-stream')
            
            attachment_list = [{
                "fileName": os.path.basename(url),
                "url": url,
                "fileType": file_type
            }]
        
        try:
            result = update_progress_api(issue_id, circulation_id, opinion, token, attachment_list)
            print(f"更新结果: {result}")
            if result.get('success') == True:
                print(f"上传单号 {order_no} 成功！")
                success_count += 1
                if url:
                    sheet.cell(row=row_idx, column=remark_col + 1, value='上传成功')
                else:
                    sheet.cell(row=row_idx, column=remark_col + 1, value='上传成功（无附件）')
            else:
                error_msg = result.get('message', '未知错误')
                print(f"上传单号 {order_no} 失败: {error_msg}")
                failure_count += 1
                failed_order_nos.append(order_no)
                sheet.cell(row=row_idx, column=remark_col + 1, value=f'上传失败: {error_msg}')
        except Exception as e:
            error_msg = str(e)
            if "权限不足" in error_msg or "401" in error_msg or "unauthorized" in error_msg.lower():
                token_result = get_new_token_dialog(token)
                if token_result["cancelled"]:
                    print("用户取消输入令牌")
                    failure_count += 1
                    failed_order_nos.append(order_no)
                    sheet.cell(row=row_idx, column=remark_col + 1, value='用户取消输入令牌')
                    continue
                
                token = token_result["token"]
                config = load_config()
                config['token'] = token
                save_config(config)
                
                result = update_progress_api(issue_id, circulation_id, opinion, token, attachment_list)
                print(f"使用新令牌更新成功: {result}")
                success_count += 1
                if url:
                    sheet.cell(row=row_idx, column=remark_col + 1, value='上传成功')
                else:
                    sheet.cell(row=row_idx, column=remark_col + 1, value='上传成功（无附件）')
            else:
                print(f"处理第 {current_row} 行失败: {e}")
                failure_count += 1
                failed_order_nos.append(order_no)
                sheet.cell(row=row_idx, column=remark_col + 1, value=f'上传失败: {str(e)}')
                continue
    
    if update_progress_callback:
        update_progress_callback(100, 100, "步骤4/4: 上传研判记录完成")
    
    try:
        wb.save(judgment_excel_path)
        print("已将上传备注写入研判记录表")
    except PermissionError:
        if tk._default_root:
            messagebox.showerror("权限错误", "请关闭研判记录表Excel文件，然后重试。")
        else:
            temp_root = tk.Tk()
            temp_root.withdraw()
            messagebox.showerror("权限错误", "请关闭研判记录表Excel文件，然后重试。")
            temp_root.destroy()
    
    print("处理完成")
    
    return token, success_count, failure_count, failed_order_nos

class ProgressWindow:
    def __init__(self, parent=None):
        self.progress_window = tk.Toplevel(parent)
        self.progress_window.title("处理进度")
        self.progress_window.geometry("500x150")
        self.progress_window.resizable(False, False)
        self.progress_window.grab_set()
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.progress_window,
            variable=self.progress_var,
            maximum=100,
            length=450,
            mode='determinate'
        )
        self.progress_bar.pack(pady=20, padx=20)
        
        self.status_label = tk.Label(self.progress_window, text="准备开始...", font=("微软雅黑", 11))
        self.status_label.pack(pady=10)
        
        self.step_label = tk.Label(self.progress_window, text="", font=("微软雅黑", 9), fg="#666")
        self.step_label.pack(pady=5)
        
        self.progress_window.update_idletasks()
        width = self.progress_window.winfo_width()
        height = self.progress_window.winfo_height()
        x = (self.progress_window.winfo_screenwidth() // 2) - (width // 2)
        y = (self.progress_window.winfo_screenheight() // 2) - (height // 2)
        self.progress_window.geometry(f'{width}x{height}+{x}+{y}')
    
    def update(self, current, total, status):
        try:
            percent = (current / total) * 100 if total > 0 else 0
            self.progress_var.set(percent)
            self.status_label.config(text=status)
            self.progress_window.update()
        except tk.TclError:
            pass
    
    def close(self):
        try:
            self.progress_window.destroy()
        except tk.TclError:
            pass

def main():
    config = load_config()
    saved_token = config.get('token', '')
    
    root = tk.Tk()
    root.title("矛盾纠纷数据处理系统")
    root.geometry("700x400")
    root.resizable(False, False)
    
    judgment_excel_path_var = tk.StringVar()
    save_path_var = tk.StringVar(value=config.get('default_save_path', DEFAULT_SAVE_PATH))
    token_var = tk.StringVar(value=saved_token)
    judgment_type_var = tk.StringVar(value="一级研判")
    
    result = {'judgment_excel': '', 'save_path': '', 'token': '', 'judgment_type': '', 'cancelled': True}
    
    def browse_judgment_excel():
        file_path = filedialog.askopenfilename(
            title="选择研判记录表",
            filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        if file_path:
            judgment_excel_path_var.set(file_path)
    
    def browse_save_path():
        selected = filedialog.askdirectory(title="选择保存路径", initialdir=save_path_var.get())
        if selected:
            save_path_var.set(selected)
    
    def confirm():
        result['judgment_excel'] = judgment_excel_path_var.get()
        result['save_path'] = save_path_var.get()
        result['token'] = token_var.get()
        result['judgment_type'] = judgment_type_var.get()
        result['cancelled'] = False
        
        if result['token']:
            config['token'] = result['token']
        if result['save_path']:
            config['default_save_path'] = result['save_path']
        save_config(config)
        
        root.destroy()
    
    def cancel():
        result['cancelled'] = True
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", cancel)
    
    tk.Label(root, text="研判记录表：", font=("微软雅黑", 10), width=20, anchor=tk.W).place(x=20, y=30)
    tk.Entry(root, textvariable=judgment_excel_path_var, width=45).place(x=160, y=30)
    tk.Button(root, text="浏览...", command=browse_judgment_excel, relief=tk.FLAT, bg="#f0f0f0", padx=10).place(x=580, y=27)
    
    tk.Label(root, text="Word保存路径：", font=("微软雅黑", 10), width=20, anchor=tk.W).place(x=20, y=70)
    tk.Entry(root, textvariable=save_path_var, width=45).place(x=160, y=70)
    tk.Button(root, text="浏览...", command=browse_save_path, relief=tk.FLAT, bg="#f0f0f0", padx=10).place(x=580, y=67)
    
    tk.Label(root, text="用户令牌：", font=("微软雅黑", 10), width=20, anchor=tk.W).place(x=20, y=110)
    tk.Entry(root, textvariable=token_var, width=55).place(x=160, y=110)
    
    tk.Label(root, text="研判类型：", font=("微软雅黑", 10), width=20, anchor=tk.W).place(x=20, y=150)
    tk.Radiobutton(root, text="一级研判", variable=judgment_type_var, value="一级研判").place(x=160, y=150)
    tk.Radiobutton(root, text="二级研判", variable=judgment_type_var, value="二级研判").place(x=260, y=150)
    tk.Radiobutton(root, text="周研判", variable=judgment_type_var, value="周研判").place(x=360, y=150)
    
    tk.Label(root, text="处理流程：", font=("微软雅黑", 10, "bold"), width=20, anchor=tk.W).place(x=20, y=200)
    steps = [
        "1. 从API获取矛盾纠纷数据",
        "2. 生成研判报告Word并上传",
        "3. 根据单号查找id和circulationId",
        "4. 上传研判记录到API"
    ]
    for i, step in enumerate(steps):
        tk.Label(root, text=step, font=("微软雅黑", 9), fg="#666").place(x=40, y=225 + i*25)
    
    tk.Button(root, text="开始处理", command=confirm, width=15, bg="#4CAF50", fg="white", relief=tk.FLAT, padx=15, pady=5).place(x=220, y=350)
    tk.Button(root, text="取消", command=cancel, width=15, relief=tk.FLAT, bg="#f0f0f0", padx=15, pady=5).place(x=380, y=350)
    
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')
    
    root.mainloop()
    
    if result['cancelled']:
        print("操作取消")
        return
    
    judgment_excel = result['judgment_excel']
    save_path = result['save_path']
    token = result['token']
    judgment_type = result['judgment_type']
    
    if not judgment_excel:
        messagebox.showerror("错误", "请选择研判记录表")
        return
    
    if not token:
        messagebox.showerror("错误", "请输入用户令牌")
        return
    
    progress_window = ProgressWindow()
    
    try:
        print("步骤1：从API获取矛盾纠纷数据...")
        fetch_dispute_data(token, progress_window.update)
        
        print("步骤2：生成研判报告Word并上传...")
        token = generate_judgment_reports(judgment_excel, save_path, judgment_type, token, progress_window.update)
        
        print("步骤3：查找id和circulationId并上传研判记录...")
        token, success_count, failure_count, failed_order_nos = process_update_progress(judgment_excel, token, progress_window.update)
        
        progress_window.update(100, 100, "所有处理已完成！")
        
        import time
        time.sleep(0.5)
        progress_window.close()
        
        total_count = success_count + failure_count
        message = f"处理完成！\n\n总处理条数: {total_count}\n成功条数: {success_count}\n失败条数: {failure_count}"
        
        if failed_order_nos:
            failed_str = "\n".join(failed_order_nos)
            message += f"\n\n未成功的单号:\n{failed_str}"
        
        messagebox.showinfo("处理完成", message)
        
    except Exception as e:
        print(f"处理失败: {e}")
        try:
            progress_window.close()
        except:
            pass
        messagebox.showerror("错误", f"处理失败: {str(e)}")

if __name__ == "__main__":
    main()